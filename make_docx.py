"""
SANJAYA EDGE — Detailed Submission Document Generator
Produces a professional Word doc with code, packages, and assumptions.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
OUT  = os.path.join(BASE, "SANJAYA_EDGE_Detailed_Document.docx")

# ─── Helpers ──────────────────────────────────────────────────────────────────
def set_font(run, size=11, bold=False, italic=False,
             color=(0,0,0), font_name="Calibri"):
    run.font.size     = Pt(size)
    run.font.bold     = bold
    run.font.italic   = italic
    run.font.name     = font_name
    run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    sizes = {1: 18, 2: 14, 3: 12}
    colors = {1: (0x3B,0x82,0xF6), 2: (0x1E,0x40,0xAF), 3: (0x1E,0x3A,0x8A)}
    set_font(run, size=sizes.get(level,12), bold=True,
             color=colors.get(level,(0,0,0)))
    # Bottom border for h1
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '3B82F6')
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p

def add_para(doc, text, size=11, bold=False, italic=False,
             color=(30,30,30), before=2, after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic, color=color)
    return p

def add_bullet(doc, text, sub=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    set_font(run, size=10.5, bold=False, color=(30,30,30))
    if sub:
        p2 = doc.add_paragraph(style="List Bullet 2")
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(2)
        r2 = p2.add_run(sub)
        set_font(r2, size=9.5, italic=True, color=(80,100,130))

def add_code_block(doc, code_text, caption=None):
    if caption:
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(6)
        cp.paragraph_format.space_after  = Pt(1)
        r = cp.add_run(caption)
        set_font(r, size=10, bold=True, color=(30,80,160))

    for line in code_text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Inches(0.3)
        # Light grey shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'F0F4FF')
        pPr.append(shd)
        run = p.add_run(line if line else ' ')
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x3A)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def read_file(path):
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        return f"[Error reading file: {e}]"

def add_table_row(table, col1, col2, header=False):
    row = table.add_row()
    for i, txt in enumerate([col1, col2]):
        cell = row.cells[i]
        cell.text = txt
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.bold = header
                run.font.name = "Calibri"
                if header:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    if header:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '1E3A8A')
            tcPr.append(shd)


# ─── Build document ───────────────────────────────────────────────────────────
def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # ── Cover ────────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(36)
    r = title_p.add_run("SANJAYA EDGE")
    set_font(r, size=32, bold=True, color=(0x3B,0x82,0xF6), font_name="Calibri")

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub_p.add_run("AI Road Safety Co-Pilot")
    set_font(r2, size=18, bold=False, color=(0x1E,0x40,0xAF))

    tag_p = doc.add_paragraph()
    tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = tag_p.add_run(
        "IIT Madras CoERS Road Safety Hackathon 2026  |  DriveLegal Track"
    )
    set_font(r3, size=12, italic=True, color=(80,100,130))

    doc.add_paragraph()
    detail_p = doc.add_paragraph()
    detail_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = detail_p.add_run(
        "Detailed Submission Document\n"
        "Includes: Solution Description  |  Complete Source Code  |"
        "  Package List  |  Assumptions"
    )
    set_font(r4, size=10.5, color=(80,100,130))

    doc.add_page_break()

    # ── 1. Project Overview ──────────────────────────────────────────────────
    add_heading(doc, "1.  Project Overview", 1)
    add_para(doc,
        "Sanjaya Edge is a real-time, citizen-facing AI road safety co-pilot. "
        "It analyses video from any camera (dashcam, CCTV, traffic pole), detects "
        "five categories of traffic violations simultaneously, and for each detection "
        "immediately delivers: the exact Motor Vehicles Act / IPC section violated, "
        "the national fine, the state-specific fine (Tamil Nadu, Maharashtra, Karnataka), "
        "a risk explanation, and a spoken voice warning in Indian English.",
        size=11)

    add_para(doc,
        "Unlike existing enforcement systems — which are reactive, police-operated, "
        "and legally opaque — Sanjaya Edge is proactive and citizen-facing. "
        "It turns every dashcam into a legal co-pilot.",
        size=11, italic=True, color=(0x1E,0x40,0xAF))

    # ── 2. Problem Statement ─────────────────────────────────────────────────
    add_heading(doc, "2.  Problem Statement", 1)
    add_para(doc,
        "India recorded 1,68,491 road crash deaths in 2022 (MoRTH Annual Report). "
        "53% of fatalities are in the 18–45 working-age group. The top three causes "
        "are: red-light jumping, wrong-side driving, and riding without a helmet. "
        "Fewer than 2% of violations ever result in a challan. The root causes are:",
        size=11)
    for item in [
        "Reactive enforcement: police act after the fact, not before harm occurs.",
        "Legal opacity: challans cite section numbers; citizens do not know what they broke.",
        "Low camera density: most violations go entirely undetected.",
        "Single-class tools: existing products detect one violation type at a time.",
    ]:
        add_bullet(doc, item)

    # ── 3. Solution Description ──────────────────────────────────────────────
    add_heading(doc, "3.  Solution Description", 1)
    add_para(doc,
        "Sanjaya Edge runs five independent detectors simultaneously on every video frame, "
        "using a single YOLOv8n pass as the shared perception backbone:", size=11)

    detectors = [
        ("Detector 1 — Red Light Jump (MVA Section 119)",
         "YOLOv8n detects the traffic light bounding box (COCO class 9). "
         "The ROI is converted to HSV colour space and split into top/mid/bottom thirds. "
         "Red pixel density and brightness position determine the colour. "
         "A 4-frame temporal consensus filter eliminates single-frame false positives "
         "caused by glare or sensor noise. When a vehicle is detected in the "
         "intersection zone while the light is confirmed red, a CRITICAL alert fires. "
         "Fine: Rs. 1,000 (first offence) / Rs. 2,000 (repeat)."),
        ("Detector 2 — Helmet Violation (MVA Section 129)",
         "A custom YOLOv8 head/helmet classifier runs on rider crops. Classes: "
         "head (bare), helmet, person. If a bare head (class: head) is detected "
         "with confidence > 38%, a HIGH RISK alert fires. "
         "Fallback: if the custom model is absent, any motorcycle detection "
         "triggers a compliance-unverified warning. "
         "Fine: Rs. 1,000 + 3-month licence suspension."),
        ("Detector 3 — Wrong-Side Driving (MVA Section 184)",
         "Vehicle centres are extracted from YOLO detections each frame. "
         "Lucas-Kanade sparse optical flow (OpenCV calcOpticalFlowPyrLK) tracks "
         "movement vectors between consecutive frames. The dominant traffic "
         "direction is computed as the circular mean of all velocity angles over "
         "a 20-frame history. Any vehicle moving more than 140 degrees away from "
         "the dominant direction is flagged as wrong-side. "
         "Fine: Rs. 5,000 (dangerous driving)."),
        ("Detector 4 — Traffic Blocking / Gridlock (IPC Section 283 / MVA Section 184)",
         "All vehicle bounding boxes are collected each frame. Vectorised pairwise "
         "IoU (padded boxes) checks which vehicles are densely packed. When 6 or "
         "more vehicles form a cluster, a CRITICAL alert fires with confidence "
         "= 58 + (vehicle_count * 4.5), capped at 95%. "
         "Fine: Rs. 5,000 + Rs. 2,000 (public nuisance)."),
        ("Detector 5 — Pedestrian in Active Lane (MVA Section 283)",
         "All person class detections with confidence > 38% are filtered by "
         "two criteria: (1) aspect ratio > 1.2 (standing pedestrian, not seated "
         "rider), and (2) centre point within the central traffic lane zone "
         "(not footpath edges). A HIGH RISK alert fires on positive detection. "
         "Fine: Rs. 500 (obstructing traffic)."),
    ]
    for (title, desc) in detectors:
        add_heading(doc, title, 2)
        add_para(doc, desc, size=10.5)

    # ── 4. Architecture ──────────────────────────────────────────────────────
    add_heading(doc, "4.  System Architecture", 1)
    add_para(doc,
        "The system is split into two services communicating over WebSocket:", size=11)

    add_heading(doc, "Backend  (Python · FastAPI · Uvicorn)", 2)
    for item in [
        "FastAPI WebSocket endpoint streams annotated frames + JSON alerts to the frontend.",
        "YOLOv8n inference runs in a concurrent.futures.ThreadPoolExecutor (2 workers) "
        "so the async event loop is never blocked.",
        "Input frames are resized to 416px wide (imgsz=320) before inference — "
        "approximately 4x faster than full-resolution with negligible accuracy loss.",
        "Frame skip = 3: YOLO runs every third frame; results are reused between runs.",
        "CUDA is auto-detected via torch.cuda.is_available(); falls back to CPU.",
        "Model is warmed up at startup with a dummy forward pass to eliminate cold-start lag.",
        "rules.json holds all violation metadata: law section, fines, risk, severity.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Frontend  (React 19 · Vite · Tailwind CSS)", 2)
    for item in [
        "WebSocket client receives base64-encoded JPEG frames and renders them as a live feed.",
        "Compliance event log displays color-coded cards: red = CRITICAL, amber = HIGH, green = SAFE.",
        "5-metric stats strip: Total Events, Critical, High Risk, Avg Confidence, Fine Exposure.",
        "Voice co-pilot: Web Speech API (lang='en-IN') reads each violation and fine on detection.",
        "HUD overlay: location, GPS, speed, detector status, REC indicator, frame counter.",
    ]:
        add_bullet(doc, item)

    # ── 5. Packages ──────────────────────────────────────────────────────────
    add_heading(doc, "5.  Software Packages Used", 1)

    add_heading(doc, "Python (Backend)", 2)
    py_packages = [
        ("fastapi", "0.111.0",    "Async web framework for WebSocket streaming"),
        ("uvicorn[standard]", "0.29.0", "ASGI server"),
        ("ultralytics", "8.2.18", "YOLOv8 model inference"),
        ("opencv-python", "4.9.0.80", "Image processing: HSV, optical flow, encoding"),
        ("numpy", "1.26.4",       "Numerical operations for vectorised IoU"),
        ("websockets", "12.0",    "WebSocket protocol support"),
        ("python-multipart", "0.0.9", "Form data parsing"),
        ("torch / torchvision", "auto", "PyTorch backend for YOLOv8 (CUDA if available)"),
        ("python-pptx", "1.0.2",  "Presentation generation (dev utility)"),
        ("python-docx", "latest", "Document generation (dev utility)"),
    ]
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    add_table_row(tbl, "Package", "Version", header=True)
    add_table_row(tbl, "", "")  # fix: remove blank row added by helper
    hrow = tbl.rows[0]
    for i, txt in enumerate(["Package", "Version", "Purpose"]):
        hrow.cells[i].text = txt
        for r in hrow.cells[i].paragraphs[0].runs:
            r.font.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = hrow.cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1E3A8A')
        tcPr.append(shd)
    tbl.rows[1]._tr.getparent().remove(tbl.rows[1]._tr)

    for (pkg, ver, purpose) in py_packages:
        row = tbl.add_row()
        for i, txt in enumerate([pkg, ver, purpose]):
            row.cells[i].text = txt
            for r in row.cells[i].paragraphs[0].runs:
                r.font.size = Pt(10)
                r.font.name = "Calibri"

    doc.add_paragraph()

    add_heading(doc, "JavaScript / Node (Frontend)", 2)
    js_packages = [
        ("react", "^19.0.0",          "UI framework"),
        ("react-dom", "^19.0.0",       "DOM rendering"),
        ("vite", "^8.0.0",             "Build tool and dev server"),
        ("@vitejs/plugin-react", "^4.5.2", "React HMR support"),
        ("tailwindcss", "^3.x",        "Utility-first CSS framework"),
        ("eslint", "^9.x",             "Code linting"),
        ("Web Speech API", "browser",  "Voice synthesis (built-in browser API)"),
    ]
    tbl2 = doc.add_table(rows=1, cols=3)
    tbl2.style = "Table Grid"
    hrow2 = tbl2.rows[0]
    for i, txt in enumerate(["Package", "Version", "Purpose"]):
        hrow2.cells[i].text = txt
        for r in hrow2.cells[i].paragraphs[0].runs:
            r.font.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = hrow2.cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1E3A8A'); tcPr.append(shd)
    for (pkg, ver, purpose) in js_packages:
        row = tbl2.add_row()
        for i, txt in enumerate([pkg, ver, purpose]):
            row.cells[i].text = txt
            for r in row.cells[i].paragraphs[0].runs:
                r.font.size = Pt(10); r.font.name = "Calibri"

    doc.add_paragraph()

    # ── 6. Assumptions ───────────────────────────────────────────────────────
    add_heading(doc, "6.  Assumptions", 1)
    assumptions = [
        ("Camera perspective",
         "The system assumes a forward-facing dashcam or fixed overhead/side CCTV angle. "
         "Detections may degrade with extreme angles (>45 degrees off-axis)."),
        ("Traffic light visibility",
         "The red light detector assumes the traffic light is visible in the frame and "
         "sufficiently lit. Night performance relies on CLAHE enhancement."),
        ("Helmet model availability",
         "The custom helmet classifier (helmet_model.pt) is an optional enhancement. "
         "If absent, the system falls back to motorcycle proxy detection, which flags "
         "all motorcyclists as unverified (not confirmed bare-head)."),
        ("Single direction of travel",
         "The wrong-side detector assumes most vehicles in the frame are travelling "
         "in one dominant direction. Wide-angle shots of junctions with multi-direction "
         "flow may produce false positives."),
        ("Indian traffic context",
         "The legal database covers MVA 1988 and IPC 283. Fine amounts reflect "
         "2023-24 values. State fines cover Tamil Nadu, Maharashtra, Karnataka only."),
        ("Video quality",
         "Minimum recommended resolution is 480p. Lower resolutions reduce YOLO "
         "detection confidence, particularly for small or distant objects."),
        ("Real-time definition",
         "Real-time means one annotated frame delivered per source video frame "
         "(at source FPS). Processing time is bounded by YOLO inference latency "
         "(typically 20-80ms per frame on CPU; <10ms on CUDA)."),
        ("Browser requirement",
         "Voice co-pilot requires a browser supporting Web Speech API "
         "(Chrome, Edge, Safari). Firefox partial support."),
        ("Port availability",
         "Backend assumes port 8000 is free. Frontend dev server uses port 5173."),
        ("Python version",
         "Tested on Python 3.11. May work on 3.10 but not guaranteed on 3.9 or below."),
    ]
    for (title, desc) in assumptions:
        add_bullet(doc, title, sub=desc)

    # ── 7. Complete Source Code ───────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "7.  Complete Source Code", 1)
    add_para(doc,
        "All source files are included below. The full repository is also available at: "
        "https://github.com/Rajbharti06/SANJAYA-EDGE",
        size=11, italic=True, color=(0x1E,0x40,0xAF))

    files = [
        ("backend/main.py",      "Backend — FastAPI Server + All 5 Detectors"),
        ("backend/rules.json",   "Legal Database — MVA 1988 + IPC 283"),
        ("backend/requirements.txt", "Python Dependencies"),
        ("frontend/src/App.jsx", "Frontend — React Dashboard"),
        ("frontend/src/index.css", "Frontend — CSS Design System"),
        ("frontend/src/main.jsx", "Frontend — React Entry Point"),
        ("frontend/index.html",  "Frontend — HTML Shell"),
        ("frontend/vite.config.js", "Frontend — Vite Configuration"),
        ("frontend/tailwind.config.js", "Frontend — Tailwind Configuration"),
    ]

    for (rel_path, caption) in files:
        full_path = os.path.join(BASE, rel_path.replace("/", os.sep))
        content   = read_file(full_path)
        add_heading(doc, rel_path, 2)
        add_para(doc, caption, size=10, italic=True, color=(80,100,130), before=0)
        add_code_block(doc, content)

    # ── 8. How to Run ─────────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "8.  How to Run the Project", 1)

    add_heading(doc, "Step 1 — Backend", 2)
    add_code_block(doc,
        "cd backend\n"
        "pip install -r requirements.txt\n"
        "uvicorn main:app --reload\n"
        "# Server starts at http://localhost:8000\n"
        "# Health check: GET http://localhost:8000/health",
        caption="Terminal commands:")

    add_heading(doc, "Step 2 — Frontend", 2)
    add_code_block(doc,
        "cd frontend\n"
        "npm install\n"
        "npm run dev\n"
        "# App starts at http://localhost:5173",
        caption="Terminal commands:")

    add_heading(doc, "Step 3 — Use", 2)
    for item in [
        "Open http://localhost:5173 in Chrome or Edge.",
        "Select a scenario from the top bar (Red Light / Helmet / Traffic / Wrong-Side).",
        "Press the blue 'START SCAN' button.",
        "The backend streams annotated frames via WebSocket.",
        "Violations appear in the event log on the right with law section and fine.",
        "The voice co-pilot reads each violation aloud in Indian English.",
    ]:
        add_bullet(doc, item)

    # ── Footer ───────────────────────────────────────────────────────────────
    doc.add_paragraph()
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run(
        "SANJAYA EDGE  |  IIT Madras CoERS Road Safety Hackathon 2026  |  "
        "github.com/Rajbharti06/SANJAYA-EDGE"
    )
    set_font(r, size=9, color=(120,140,160), italic=True)

    doc.save(OUT)
    sz = os.path.getsize(OUT) / (1024 * 1024)
    print(f"[docx] Saved: {OUT}  ({sz:.1f} MB)")


if __name__ == "__main__":
    main()
